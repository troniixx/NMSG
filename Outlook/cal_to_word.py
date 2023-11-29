from ics import Calendar
from docx import Document
from datetime import datetime, timedelta

def parse_ics(path):
    with open(path, "r") as cal:
        calendar = Calendar(cal.read())
    
    return calendar

def create_word(events, filename):
    doc = Document()

    for event in events:
        doc.add_heading(event.name, level = 1)
        doc.add_paragraph(event.description)
        doc.add_paragraph(event.location)
        doc.add_paragraph("Start: " + str(event.begin.datetime))
        doc.add_paragraph("End: " + str(event.end.datetime))
        doc.add_paragraph("-------------")
    
    doc.save(filename)

if __name__ == "__main__":
    
    ics_path = ""
    cal = parse_ics(ics_path)

    start_date = datetime(2023, 12, 1)
    end_date = start_date + timedelta(days = 7)

    events_week = [event for event in cal.events if start_date <= event.begin.datetime < end_date]

    create_word(events_week, "calendar_weekly.docx")